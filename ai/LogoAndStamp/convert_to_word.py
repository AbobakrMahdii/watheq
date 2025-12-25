"""Convert Markdown to Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re


def md_to_docx(md_path, docx_path):
    """Convert markdown file to Word document."""
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code block
        if line.startswith('```'):
            if in_code_block:
                # End code block - add as formatted paragraph
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Table
        if line.startswith('|') and '|' in line[1:]:
            table_rows.append(line)
            in_table = True
            i += 1
            continue
        elif in_table and not line.startswith('|'):
            # Process table
            add_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=3)
        elif line.startswith('---'):
            # Horizontal rule - add empty paragraph
            doc.add_paragraph()
        elif line.strip() == '':
            pass  # Skip empty lines
        elif not in_table:
            # Regular paragraph
            p = doc.add_paragraph()
            add_formatted_text(p, line)
        
        i += 1
    
    # Handle remaining table
    if table_rows:
        add_table(doc, table_rows)
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def add_table(doc, rows):
    """Add a table to the document."""
    if len(rows) < 2:
        return
    
    # Parse rows
    data = []
    for row in rows:
        if '---' in row:
            continue  # Skip separator row
        cells = [c.strip() for c in row.split('|')[1:-1]]
        if cells:
            data.append(cells)
    
    if not data:
        return
    
    # Create table
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = clean_text(cell_text)
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    doc.add_paragraph()  # Space after table


def add_formatted_text(paragraph, text):
    """Add formatted text to paragraph."""
    text = clean_text(text)
    
    # Handle bold (**text**)
    parts = re.split(r'\*\*(.+?)\*\*', text)
    is_bold = False
    for part in parts:
        if part:
            run = paragraph.add_run(part)
            run.bold = is_bold
        is_bold = not is_bold


def clean_text(text):
    """Remove markdown formatting."""
    # Remove inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove bold markers for plain text
    text = text.replace('**', '')
    return text.strip()


if __name__ == '__main__':
    md_to_docx(
        'docs/Chapter4_Increment1_LogoVerification.md',
        'docs/Chapter4_Increment1_LogoVerification.docx'
    )
